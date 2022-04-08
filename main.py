import tkinter as tk
from tkinter import ttk, filedialog
import pandas as pd
import sqlite3
from tkinter import messagebox
from docx import Document
from docx.shared import Pt

class Application:
    def __init__(self):
        self.window = tk.Tk()
        self.window.configure(bg='#ffa366')
        self.window.title("piekarnia")
        self.window.iconbitmap('C:/bakery_app_files/logo3.ico')
        self.fonts = ("Comic Sans MS", 9, "bold")
        self.app_width = self.window.winfo_screenwidth()
        self.app_height = self.window.winfo_screenheight()
        self.scr_width = self.window.winfo_screenwidth()
        self.scr_height = self.window.winfo_screenheight()
        self.x = (self.scr_width / 2) - (self.app_width / 2)
        self.y = (self.scr_height / 2) - (self.app_height / 2)
        self.window.geometry(f'{self.app_width}x{self.app_height}+{int(self.x)}+{int(self.y)}')
        self.notebook = ttk.Notebook(self.window)
        self.notebook.pack(pady=25)

        self.tab1 = tk.Frame(self.notebook, width=self.app_width - 50, height=self.app_height - 150, bg='#F0F0F0')
        self.tab1.pack(fill="both", expand=True, ipadx=15, ipady=15)
        self.notebook.add(self.tab1, text="produkcja")

        self.tab2 = tk.Frame(self.notebook, width=self.app_width - 50, height=self.app_height - 150)
        self.tab2.pack(fill="both", expand=True, ipadx=15, ipady=15)
        self.notebook.add(self.tab2, text="   bary   ")
        self.mainframe=ttk.LabelFrame(self.tab1,text="menu główne",width=self.app_width-150)
        self.mainframe.pack(expand="yes",fill="both", side=tk.TOP)
        self.bottomframe=ttk.LabelFrame(self.tab1, text="tabela", width=self.app_width-150,height=self.app_height-340)
        self.bottomframe.pack(expand="yes", fill="both")
        self.buttons_frame = ttk.LabelFrame(self.tab1, text="opcje", width=self.app_width-150,)
        self.buttons_frame.pack(expand="yes", fill="both",side=tk.BOTTOM)
        ''' ==========================PRZYCISKI MENU ==================================================='''
        self.select_button = tk.Button(self.buttons_frame, text="wyczyść pozycję", command=self.clear,state=tk.DISABLED,bd=3,
                                       width=15, height=2, activebackground='#ffa366', relief="groove")
        self.select_button.grid(row=0, column=0, padx=12, pady=5)

        self.update_button = tk.Button(self.buttons_frame, text="zapisz edycję",command=self.update_record,state=tk.DISABLED, bd=3,
                                       width=15, height=2, activebackground='#ffa366',relief="groove")
        self.update_button.grid(row=0, column=1, padx=12, pady=5)

        self.remove_button = tk.Button(self.buttons_frame, text="usuń pozycję", command=self.remove,state=tk.DISABLED,bd=3,
                                       width=15, height=2, activebackground='#ffa366',relief="groove")
        self.remove_button.grid(row=0, column=2, padx=12, pady=5)

        self.add_button = tk.Button(self.buttons_frame, text="dodaj pozycję", state=tk.DISABLED, bd=3, width=15, height=2,
                                    activebackground='#ffa366',relief="groove")
        self.add_button.grid(row=0, column=3, padx=12, pady=5)
        ''' ==========================PRZYCISKI MENU GŁÓWNEGO==================================================='''
        self.button1 = tk.Button(self.mainframe, text="wczytaj plik", command=self.open_file, bd=3, width=15, height=2,
                                 activebackground='#ffa366',relief="groove")
        self.button1.grid(column=1, row=1, padx=15, pady=15)
        self.button2 = tk.Button(self.mainframe, text="otwórz zestawienie", command=self.create_view, bd=3, width=15,
                                 height=2, activebackground='#ffa366',relief="groove",state=tk.DISABLED)
        self.button2.grid(column=2, row=1, padx=15, pady=15)
        self.button3 = tk.Button(self.mainframe, text="zapisz",  bd=3, width=15, height=2, activebackground='#ffa366'
                                 , relief="groove",state=tk.DISABLED)
        self.button3.grid(column=3, row=1, padx=15, pady=15)
        self.button4 = tk.Button(self.mainframe, text="drukuj", bd=3, width=15, height=2, activebackground='#ffa366',
                                 relief="groove",state=tk.DISABLED,command=self.print_word)
        self.button4.grid(column=4, row=1, padx=15, pady=15)
        #self.info_label = tk.Label(self.tab1, text="")
        #self.info_label.grid(row=2, column=1)

        self.style = ttk.Style()
        #self.style.theme_use('default')
        self.style.configure('TLabel', foreground='red', background='grey')
        self.style.configure('TButton', foreground='white', background='blue')
        self.style.configure('TFrame', background='black')

        self.clicked = tk.StringVar()
        self.clicked.set("Wybierz dzień")
        self.drop = tk.OptionMenu(self.mainframe, self.clicked, "Poniedziałek", "Wtorek", "Środa", "Czwartek", "Piątek",
                                  "Sobota",command=self.choose)
        #self.drop.config(font=self.fonts)
        self.drop.grid(column=2, row=0)
    def choose(self,clicked):
        if self.clicked.get() != "Wybierz dzień":
            if (self.button2['state'] == tk.DISABLED):
                self.button2['state'] = tk.ACTIVE
            else:
                self.button2['state'] = tk.NORMAL
        else:
            return
        if self.clicked.get() != "Wybierz dzień":
            if (self.button3['state'] == tk.DISABLED):
                self.button3['state'] = tk.ACTIVE
            else:
                self.button3['state'] = tk.NORMAL
        else:
            return
        if self.clicked.get() != "Wybierz dzień":
            if (self.button4['state'] == tk.DISABLED):
                self.button4['state'] = tk.ACTIVE
            else:
                self.button4['state'] = tk.NORMAL
        else:
            return

    def open_file(self):
        filepath = filedialog.askopenfilename(initialdir="C:\\Users\\user\\PycharmProjects\\bakery_app",
                                              title="wybierz plik z dzisiejszym planem produkcji",
                                              filetypes=(("xls files", "*.xls"), ("All files", "*.*")))

        if filepath:
            try:
                filepath = r"{}".format(filepath)
                df = pd.read_excel(filepath, index_col=0, usecols=[4, 11], skiprows=[i for i in range(1, 8)])
                df.dropna(axis=0, inplace=True)
                self.idx = df.index
            except ValueError:
                messagebox.showerror("BŁĄD", "nie można znaleźć pliku, spróbuj ponownie")
            except FileNotFoundError:
                messagebox.showerror("BŁĄD","nie można znaleźć pliku, spróbuj ponownie")

        self.values = []
        for i in range(0, len(df)):
            self.values.append(df.iat[i, 0])

        conn = sqlite3.connect('costumer.db')
        c = conn.cursor()
        c.execute(("""DROP TABLE IF EXISTS produkcja_sklepy"""))
        c.execute("""CREATE TABLE produkcja_sklepy (
        id_produkt text primary key,
        ilosc integer
        )""")

        for i in range(0, len(df)):
            c.execute("""INSERT INTO produkcja_sklepy VALUES (?,?)""", (self.idx[i], self.values[i]))

        c.execute("""SELECT * FROM produkcja_sklepy""")
        conn.commit()
        conn.close()

    def create_view(self):

        self.options_buttons=[self.select_button,self.update_button,self.remove_button,self.add_button]
        for button in self.options_buttons:
            if (button['state'] == tk.DISABLED):
                button['state'] = tk.ACTIVE
            else:
                button['state'] = tk.NORMAL

        conn = sqlite3.connect('costumer.db')
        c = conn.cursor()
        c.execute(("""DROP TABLE IF EXISTS produkcja_teren"""))
        c.execute("""CREATE TABLE produkcja_teren (
                id_produkt text primary key,
                ilosc_teren1 integer,
                ilosc_teren2 integr
                )""")
        nazwy = ['1CHLEB 1 KG', '1CHLEB 1 KG FOR', '1CHLEB 0,70', '1CHLEB 0,70 KRO', '1CHLEB 0,60', '1CHLEB 0,60 KRO',
                 '2R INDYJSKI', '2RAZOWY 0,60',
                 '4WROCŁAWSKA', '3BUŁKA 0,10', '3BUŁKA 0,05', '5M ROGAL 0,10', '5M WARKOCZ 0,10', '5M CHAŁKA 0,50',
                 '6S BUŁ BUDYŃ', '6S BUŁ JABŁ', '6S BUŁ MAK ', '6S BUŁ SER', '6S PĄCZEK', '6S PĄCZEK DONUT',
                 '6S PĄCZEK Z SER']

        if self.clicked.get() == "Poniedziałek":
            ilosc1 = [2, 0, 20, 15, 0, 0, 0, 2, 1, 15, 15, 5, 0, 1, 0, 0, 2, 2, 2, 0, 0]
            ilosc2 = [1, 1, 1, 1, 1, 1, 1, 1, 1, 10, 20, 3, 2, 1, 2, 2, 2, 2, 2, 0, 0]
        elif self.clicked.get() == "Wtorek":
            ilosc1 = [2, 0, 20, 15, 0, 0, 0, 2, 1, 15, 15, 5, 0, 1, 0, 0, 2, 2, 2, 0, 0]
            ilosc2 = [2, 2, 2, 2, 2, 2, 2, 1, 1, 10, 20, 3, 2, 1, 2, 2, 2, 2, 2, 0, 0]
        elif self.clicked.get() == "Środa":
            ilosc1 = [2, 0, 20, 15, 0, 0, 0, 2, 1, 15, 15, 5, 0, 1, 0, 0, 2, 2, 2, 0, 0]
            ilosc2 = [3, 3, 3, 3, 3, 3, 3, 3,3, 3, 20, 3, 2, 1, 2, 2, 2, 2, 2, 0, 0]
        elif self.clicked.get() == "Czwartek":
            ilosc1 = [2, 0, 20, 15, 0, 0, 0, 2, 1, 15, 15, 5, 0, 1, 0, 0, 2, 2, 2, 0, 0]
            ilosc2 = [4, 4, 4, 4, 4, 4, 4, 1, 1, 10, 20, 3, 2, 1, 2, 2, 2, 2, 2, 0, 0]
        elif self.clicked.get() == "Piątek":
            ilosc1 = [2, 0, 20, 15, 0, 0, 0, 2, 1, 15, 15, 5, 0, 1, 0, 0, 2, 2, 2, 0, 0]
            ilosc2 = [5, 5, 5, 5, 5, 5, 5, 1, 1, 10, 20, 3, 2, 1, 2, 2, 2, 2, 2, 0, 0]
        else:
            ilosc1 = [2, 0, 20, 15, 0, 0, 0, 2, 1, 15, 15, 5, 0, 1, 0, 0, 2, 2, 2, 0, 0]
            ilosc2 = [6, 6, 6, 6, 6, 6, 6, 6, 6, 6, 20, 3, 2, 1, 2, 2, 2, 2, 2, 0, 0]
        for i in range(0, len(nazwy)):
            c.execute("""INSERT INTO produkcja_teren VALUES (?,?,?)""", (nazwy[i], ilosc1[i], ilosc2[i]))

        c.execute("""SELECT s.id_produkt,s.ilosc, t.ilosc_teren1,t.ilosc_teren2,
         COALESCE(s.ilosc + t.ilosc_teren1 +t.ilosc_teren2,s.ilosc) AS suma
        FROM produkcja_sklepy AS s
        LEFT JOIN
        produkcja_teren AS t ON s.id_produkt = t.id_produkt
        """)
        print(type(conn))
        self.records = c.fetchall()
        conn.commit()
        conn.close()

        # TREEVIEW
        self.style.configure("Treeview",
                             background="#F0F0F0",
                             foreground="black",
                             rowheight=30,
                             fieldbackground="#F0F0F0"
                             )
        self.style.map('Treeview', background=[('selected', '#c35d18')])

        self.tree_frame = ttk.Frame(self.bottomframe)
        self.tree_frame.pack()
        self.tree_scroll = ttk.Scrollbar(self.tree_frame)
        self.tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree = ttk.Treeview(self.tree_frame, yscrollcommand=self.tree_scroll.set, selectmode="extended")
        self.tree.pack()
        self.tree_scroll.config(command=self.tree.yview)
        self.tree['columns'] = ("produkt", "sklepy", "teren1", "teren2", "suma")
        self.tree.column("#0", width=0, stretch=tk.NO)
        self.tree.column("produkt", anchor=tk.W, width=190)
        self.tree.column("sklepy", anchor=tk.CENTER, width=150)
        self.tree.column("teren1", anchor=tk.CENTER, width=150)
        self.tree.column("teren2", anchor=tk.CENTER, width=150)
        self.tree.column("suma", anchor=tk.CENTER, width=150)

        self.tree.heading("#0", text="", anchor=tk.W)
        self.tree.heading("produkt", text="produkt", anchor=tk.CENTER)
        self.tree.heading("sklepy", text="sklepy", anchor=tk.CENTER)
        self.tree.heading("teren1", text="teren1", anchor=tk.CENTER)
        self.tree.heading("teren2", text="teren2", anchor=tk.CENTER)
        count = 0

        self.tree.tag_configure('oddrow', background="white")
        self.tree.tag_configure('evenrow', background="#ffa366")

        self.tree.heading("suma", text="suma", anchor=tk.CENTER)
        for i in self.tree.get_children():
            self.tree.delete(i)

        for record in self.records:
            if count % 2 == 0:
                self.tree.insert(parent='', index=tk.END, text='',
                                 values=(record[0], ('%f' % record[1]).rstrip('0').rstrip('.'), record[2], record[3],
                                         ('%f' % record[4]).rstrip('0').rstrip('.') ), tags=('evenrow',))
            else:
                self.tree.insert(parent='', index=tk.END, text='',
                                 values=(record[0], ('%f' % record[1]).rstrip('0').rstrip('.'), record[2], record[3],
                                         ('%f' % record[4]).rstrip('0').rstrip('.')), tags=('oddrow',))
            count += 1

        self.record_frame = ttk.LabelFrame(self.bottomframe, text="pozycja")
        self.record_frame.pack(expand="yes",fill="x")
        self.prod_label = tk.Label(self.record_frame, text="produkt").grid(row=0, column=0, padx=3, pady=5)
        self.prod_entry = tk.Entry(self.record_frame)
        self.prod_entry.grid(row=0, column=1, padx=3, pady=5)
        self.no_sklepy_label = tk.Label(self.record_frame, text="ilość-sklepy").grid(row=0, column=2, padx=3, pady=5)
        self.no_sklepy_entry = tk.Entry(self.record_frame)
        self.no_sklepy_entry.grid(row=0, column=3, padx=3, pady=5)
        self.no_teren1_label = tk.Label(self.record_frame, text="ilość-teren1").grid(row=0, column=4, padx=3, pady=5)
        self.no_teren1_entry = tk.Entry(self.record_frame)
        self.no_teren1_entry.grid(row=0, column=5, padx=3, pady=5)
        self.no_teren2_label = tk.Label(self.record_frame, text="ilość-teren2").grid(row=0, column=6, padx=3, pady=5)
        self.no_teren2_entry = tk.Entry(self.record_frame)
        self.no_teren2_entry.grid(row=0, column=7, padx=3, pady=5)
        self.tree.bind("<ButtonRelease-1>", self.select)

    def select(self, e):
        self.clear()
        self.selected = self.tree.focus()
        self.values = self.tree.item(self.selected, 'values')
        self.prod_entry.insert(0, self.values[0])
        self.no_sklepy_entry.insert(0, self.values[1])
        self.no_teren1_entry.insert(0, self.values[2])
        self.no_teren2_entry.insert(0, self.values[3])

    def clear(self):
        self.prod_entry.delete(0, tk.END)
        self.no_sklepy_entry.delete(0, tk.END)
        self.no_teren1_entry.delete(0, tk.END)
        self.no_teren2_entry.delete(0, tk.END)

    def remove(self):
        x = self.tree.selection()[0]
        self.tree.delete(x)

    def update_record(self):
        conn = sqlite3.connect('costumer.db')
        c = conn.cursor()

        try:
            float(self.no_sklepy_entry.get())
            if self.no_teren1_entry.get() != "None":
                float(self.no_teren1_entry.get())
            else:
                self.no_teren1_entry.delete(0,tk.END)
                self.no_teren1_entry.insert(tk.END,0)
            if self.no_teren2_entry.get() != "None":
                float(self.no_teren2_entry.get())
            else:
                self.no_teren2_entry.delete(0, tk.END)
                self.no_teren2_entry.insert(tk.END,0)

        except ValueError:
            messagebox.showerror("BŁĄD", "Błędne dane! Proszę wprowadzić wartość liczbową ")

        selected = self.tree.focus()
        self.tree.item(selected, text="", values=(self.prod_entry.get(), self.no_sklepy_entry.get(),
                                                  self.no_teren1_entry.get(), self.no_teren2_entry.get(),str(int(self.no_sklepy_entry.get())+int(self.no_teren1_entry.get())+int(self.no_teren2_entry.get()))))

        c.execute('''UPDATE produkcja_sklepy SET
                ilosc = :val
                WHERE 
                id_produkt = :id_pro''',
                    {
                        'id_pro':str(self.prod_entry.get()),
                        'val' :str(self.no_sklepy_entry.get()),
                    }
        )
        conn.commit()
        c.execute('''UPDATE produkcja_teren SET
                ilosc_teren1 = :val1,
                ilosc_teren2 = :val2
                WHERE id_produkt = :oid''',
                  {
                      'val1' : self.no_teren1_entry.get(),
                      'val2' : self.no_teren2_entry.get(),
                      'oid': self.prod_entry.get(),
                  }
        )

        conn.commit()

        conn.close()
        self.clear()

    def print_word(self):
        doc=Document()
        doc.add_heading('Produkcja na dzień:', 0)

        conn = sqlite3.connect('costumer.db')
        c = conn.cursor()
        c.execute("""SELECT s.id_produkt,
         COALESCE(s.ilosc + t.ilosc_teren1 +t.ilosc_teren2,s.ilosc) AS suma
        FROM produkcja_sklepy AS s
        LEFT JOIN
        produkcja_teren AS t ON s.id_produkt = t.id_produkt""")

        temp=c.fetchall()
        for t in range (0, len(temp)):
            nazwa=temp[t][0]
            suma= str(temp[t][1])
            dana = doc.add_paragraph().add_run(nazwa + " -- " +suma)
            dana.font.size = Pt(18)

        doc.save("produkcja.docx")

        conn.commit()


apl = Application()

apl.window.mainloop()
